
# -*- coding: utf-8 -*-
"""
生蚝AI知识库 - 去内置兜底知识库 + 扩大检索来源版

功能：
1. 不再写入任何“内置兜底知识库”，避免错误标准号污染知识库
2. 自动抓取公开网页正文
3. 自动抓取 Crossref / Semantic Scholar 开放论文题录与摘要
4. 可选通过 SerpAPI 检索：
   - 知网相关题录/摘要结果
   - 谷歌学术中文论文摘要/题录
   - 百度学术、万方、维普等公开检索结果
5. 可选解析 data/raw/ 中的 PDF、DOCX、TXT、MD、CSV、XLSX
6. 自动分段、去重、分类、摘要
7. 输出 knowledge_base.json

重要说明：
- 不自动下载知网、万方、维普、标准平台等受限/版权 PDF 全文
- 不绕过登录、验证码、付费墙
- 对于标准编号，建议前端 AI 只把带 URL 的权威来源作为依据
"""

import os
import re
import csv
import json
import time
import hashlib
import urllib.parse
from pathlib import Path
from datetime import datetime
from html import unescape

import requests
from bs4 import BeautifulSoup

RAW_DIR = Path("data/raw")
OUTPUT_FILE = Path("knowledge_base.json")

CHUNK_SIZE = 900
CHUNK_OVERLAP = 150

MAX_SEARCH_RESULTS_PER_QUERY = 8
MAX_WEB_PAGES_TOTAL = 90
MAX_SERP_RESULTS_PER_QUERY = 10
REQUEST_SLEEP = 1.2

SERPAPI_KEY = os.getenv("SERPAPI_KEY", "").strip()

HEADERS = {
    "User-Agent": "Mozilla/5.0 (compatible; OysterKnowledgeBot/3.1; +https://github.com/)"
}

# 公开网页检索关键词，进一步细化
KEYWORDS = [
    "牡蛎 养殖 技术",
    "生蚝 养殖 技术",
    "太平洋牡蛎 养殖 技术",
    "长牡蛎 养殖 技术",
    "香港牡蛎 养殖 技术",
    "近江牡蛎 养殖 技术",
    "太平洋牡蛎 种质",
    "长牡蛎 种质",
    "香港牡蛎 种质",
    "牡蛎 种质资源",
    "牡蛎 良种选育",
    "牡蛎 苗种 繁育",
    "牡蛎 人工育苗",
    "牡蛎 采苗 附着基",
    "牡蛎 筏式 吊养",
    "牡蛎 滩涂 养殖",
    "牡蛎 三倍体 育种",
    "牡蛎 四倍体 育种",
    "三倍体牡蛎 养殖",
    "牡蛎 夏季死亡",
    "牡蛎 病害 防控",
    "牡蛎 弧菌 病害",
    "牡蛎 专利 方法",
    "牡蛎 国家标准",
    "牡蛎 行业标准",
    "牡蛎 地方标准",
    "太平洋牡蛎 标准",
    "长牡蛎 标准",
    "牡蛎 论文 摘要",
]

# 学术检索关键词
SCHOLAR_QUERIES = [
    "牡蛎 种质 标准 论文 摘要",
    "太平洋牡蛎 种质资源 论文 摘要",
    "长牡蛎 种质 论文 摘要",
    "牡蛎 良种选育 论文 摘要",
    "三倍体牡蛎 育种 论文 摘要",
    "牡蛎 人工育苗 论文 摘要",
    "牡蛎 筏式养殖 论文 摘要",
    "牡蛎 病害防控 论文 摘要",
    "牡蛎 夏季死亡 论文 摘要",
    "牡蛎 弧菌 病害 论文 摘要",
]

ENGLISH_KEYWORDS = [
    "Crassostrea gigas aquaculture",
    "Pacific oyster breeding",
    "triploid oyster breeding",
    "oyster seed production",
    "oyster disease control",
    "Crassostrea gigas genetic resources",
    "Crassostrea gigas triploid",
]

PREFERRED_DOMAINS = [
    "moa.gov.cn",
    "cnfm.com.cn",
    "ysfri.ac.cn",
    "qdio.ac.cn",
    "cafs.ac.cn",
    "std.samr.gov.cn",
    "samr.gov.cn",
    "cnipa.gov.cn",
    "patents.google.com",
    "xueshu.baidu.com",
    "scholar.google.com",
    "kns.cnki.net",
    "cnki.net",
    "wanfangdata.com.cn",
    "cqvip.com",
    "semanticscholar.org",
    "crossref.org",
    "fao.org",
]

SUPPORTED_SUFFIXES = {
    ".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx"
}

CATEGORY_RULES = {
    "种质标准": [
        "种质", "标准", "GB/T", "SC/T", "地方标准", "行业标准", "品种", "良种",
        "太平洋牡蛎", "长牡蛎", "香港牡蛎", "近江牡蛎", "种质资源", "standard"
    ],
    "养殖规程": [
        "养殖", "规程", "技术规范", "筏式", "吊养", "滩涂", "苗种", "育苗",
        "采苗", "附着基", "盐度", "温度", "密度", "投饵", "水质", "敌害",
        "aquaculture", "culture"
    ],
    "专利文献": [
        "专利", "发明", "实用新型", "权利要求", "说明书", "申请号", "公开号",
        "一种", "装置", "方法", "系统", "patent"
    ],
    "学术论文": [
        "摘要", "关键词", "研究", "试验", "实验", "结果", "讨论", "参考文献",
        "DOI", "Journal", "Abstract", "article", "paper", "学位论文", "期刊"
    ],
    "遗传育种": [
        "三倍体", "四倍体", "二倍体", "育种", "家系", "选育", "杂交", "遗传",
        "分子标记", "SNP", "微卫星", "基因组", "倍性", "triploid", "breeding",
        "genetic", "genome"
    ],
    "病害防控": [
        "病害", "弧菌", "死亡", "夏季死亡", "寄生虫", "防控", "免疫", "抗病",
        "病原", "感染", "disease", "mortality", "Vibrio"
    ],
}


def safe_get(url, params=None, timeout=20):
    for i in range(3):
        try:
            resp = requests.get(url, headers=HEADERS, params=params, timeout=timeout)
            if resp.status_code == 200:
                if not resp.encoding or resp.encoding.lower() == "iso-8859-1":
                    resp.encoding = resp.apparent_encoding
                return resp
            print(f"  ⚠️ HTTP {resp.status_code}: {url}")
        except Exception as e:
            print(f"  ⚠️ 请求失败 第{i+1}次：{url}，原因：{e}")
        time.sleep(1.5)
    return None


def clean_text(text):
    if not text:
        return ""
    text = unescape(str(text))
    text = text.replace("\u3000", " ")
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    text = re.sub(r"分享到.*", "", text)
    text = re.sub(r"责任编辑[:：].*", "", text)
    text = re.sub(r"打印本页.*", "", text)
    return text.strip()


def make_id(text):
    return hashlib.md5(text.encode("utf-8", errors="ignore")).hexdigest()[:16]


def get_domain(url):
    try:
        return urllib.parse.urlparse(url).netloc
    except Exception:
        return ""


def extract_date(text):
    patterns = [
        r"\d{4}-\d{1,2}-\d{1,2}",
        r"\d{4}/\d{1,2}/\d{1,2}",
        r"\d{4}年\d{1,2}月\d{1,2}日",
        r"\b(19|20)\d{2}\b",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return m.group(0)
    return ""


def classify_text(title, content):
    full_text = f"{title} {content}"
    scores = {}
    for category, words in CATEGORY_RULES.items():
        score = 0
        for word in words:
            if word.lower() in full_text.lower():
                score += 1
        scores[category] = score
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "综合资料"


def summarize_text(text, max_len=240):
    text = clean_text(text)
    if len(text) <= max_len:
        return text

    sentences = re.split(r"[。！？；\n.!?;]", text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 10]

    keywords = [
        "目的", "方法", "结果", "结论", "规定", "要求", "适用于",
        "牡蛎", "生蚝", "太平洋牡蛎", "长牡蛎", "三倍体", "养殖",
        "育苗", "种质", "专利", "标准", "规程", "摘要",
        "oyster", "Crassostrea", "aquaculture", "breeding"
    ]

    scored = []
    for s in sentences:
        score = sum(1 for kw in keywords if kw.lower() in s.lower())
        scored.append((score, s))

    scored.sort(key=lambda x: x[0], reverse=True)
    selected = [s for score, s in scored[:3] if score > 0]
    if not selected:
        selected = sentences[:3]

    summary = "。".join(selected)
    return summary[:max_len]


def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    text = clean_text(text)
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    paragraphs = re.split(r"\n\s*\n", text)
    chunks = []
    current = ""

    for p in paragraphs:
        p = p.strip()
        if not p:
            continue

        if len(current) + len(p) + 1 <= chunk_size:
            current += "\n" + p
        else:
            if current.strip():
                chunks.append(current.strip())

            if len(p) > chunk_size:
                start = 0
                while start < len(p):
                    end = start + chunk_size
                    chunks.append(p[start:end])
                    next_start = end - overlap
                    if next_start <= start:
                        next_start = end
                    start = next_start
            else:
                tail = current[-overlap:] if len(current) > overlap else current
                current = tail + "\n" + p

    if current.strip():
        chunks.append(current.strip())

    return chunks


def create_doc(title, content, source, url="", date="", file_name="", chunk_index=1, source_type="web"):
    title = clean_text(title)[:180] or "未命名资料"
    content = clean_text(content)
    url = clean_text(url)

    has_url = bool(url.startswith("http"))

    return {
        "id": make_id(f"{title}-{source}-{url}-{chunk_index}-{content[:100]}"),
        "title": title,
        "file_name": file_name,
        "category": classify_text(title, content),
        "source": source,
        "source_type": source_type,
        "has_url": has_url,
        "content": content,
        "summary": summarize_text(content),
        "date": date,
        "url": url,
        "domain": get_domain(url) if has_url else "",
        "chunk_index": chunk_index,
        "content_length": len(content)
    }


def is_bad_url(url):
    lower = url.lower()
    bad_suffixes = [".jpg", ".jpeg", ".png", ".gif", ".zip", ".rar", ".mp4", ".mp3"]
    if any(lower.endswith(s) for s in bad_suffixes):
        return True
    if "javascript:" in lower:
        return True
    if "login" in lower or "passport" in lower:
        return True
    return False


def is_relevant_oyster(text):
    terms = [
        "oyster", "Crassostrea", "牡蛎", "生蚝", "Pacific oyster",
        "triploid oyster", "shellfish", "贝类"
    ]
    low = text.lower()
    return any(t.lower() in low for t in terms)


# ---------------------------------------------------------
# DuckDuckGo / Bing 搜索
# ---------------------------------------------------------
def search_duckduckgo(query, max_results=MAX_SEARCH_RESULTS_PER_QUERY):
    print(f"🔎 DuckDuckGo 搜索：{query}")
    urls = []
    resp = safe_get("https://duckduckgo.com/html/", params={"q": query})
    if not resp:
        return urls

    soup = BeautifulSoup(resp.text, "html.parser")
    links = soup.select("a.result__a")

    for a in links:
        href = a.get("href", "")
        if not href:
            continue

        real_url = href
        if "uddg=" in href:
            parsed = urllib.parse.urlparse(href)
            qs = urllib.parse.parse_qs(parsed.query)
            if "uddg" in qs:
                real_url = qs["uddg"][0]

        if real_url.startswith("http") and not is_bad_url(real_url):
            urls.append(real_url)

        if len(urls) >= max_results:
            break

    time.sleep(REQUEST_SLEEP)
    return urls


def search_bing(query, max_results=MAX_SEARCH_RESULTS_PER_QUERY):
    print(f"🔎 Bing 搜索：{query}")
    urls = []
    resp = safe_get("https://www.bing.com/search", params={"q": query})
    if not resp:
        return urls

    soup = BeautifulSoup(resp.text, "html.parser")
    for item in soup.select("li.b_algo h2 a"):
        href = item.get("href", "")
        if href.startswith("http") and not is_bad_url(href):
            urls.append(href)
        if len(urls) >= max_results:
            break

    time.sleep(REQUEST_SLEEP)
    return urls


def collect_web_urls():
    print("🌐 开始自动联网检索公开网页...")
    all_urls = []

    targeted_sites = (
        "site:moa.gov.cn OR site:cnfm.com.cn OR site:cafs.ac.cn OR "
        "site:qdio.ac.cn OR site:ysfri.ac.cn OR site:std.samr.gov.cn OR "
        "site:cnipa.gov.cn OR site:patents.google.com"
    )

    for kw in KEYWORDS:
        all_urls.extend(search_duckduckgo(f"{kw} {targeted_sites}"))
        all_urls.extend(search_bing(f"{kw} {targeted_sites}"))

    for kw in KEYWORDS:
        all_urls.extend(search_duckduckgo(kw))
        all_urls.extend(search_bing(kw))

    seen = set()
    unique = []
    for url in all_urls:
        normalized = url.split("#")[0]
        if normalized not in seen:
            seen.add(normalized)
            unique.append(normalized)

    def priority(u):
        score = 0
        decoded = urllib.parse.unquote(u)
        for domain in PREFERRED_DOMAINS:
            if domain in u:
                score += 10
        if "牡蛎" in decoded or "生蚝" in decoded:
            score += 5
        if "standard" in u.lower() or "patent" in u.lower():
            score += 3
        return score

    unique.sort(key=priority, reverse=True)
    selected = unique[:MAX_WEB_PAGES_TOTAL]
    print(f"  ✅ 收集到候选网页 {len(unique)} 个，选取 {len(selected)} 个")
    return selected


# ---------------------------------------------------------
# 网页正文抽取
# ---------------------------------------------------------
def extract_page_content(url):
    resp = safe_get(url)
    if not resp:
        return None

    content_type = resp.headers.get("Content-Type", "").lower()
    if "pdf" in content_type or url.lower().endswith(".pdf"):
        return None

    soup = BeautifulSoup(resp.text, "html.parser")

    for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "form"]):
        tag.decompose()

    title = ""
    if soup.title and soup.title.string:
        title = soup.title.string.strip()

    h1 = soup.find("h1")
    if h1 and h1.get_text(strip=True):
        title = h1.get_text(strip=True)

    # 尽量读取 meta description 作为补充摘要
    meta_desc = ""
    meta = soup.find("meta", attrs={"name": "description"})
    if meta and meta.get("content"):
        meta_desc = meta.get("content")

    candidates = []
    selectors = [
        "article",
        ".article",
        ".content",
        ".main",
        ".TRS_Editor",
        ".Custom_UnionStyle",
        "#content",
        "#main",
        ".detail",
        ".news_content",
        ".article-content",
        ".abstract",
        "#ChDivSummary",
    ]

    for sel in selectors:
        for node in soup.select(sel):
            text = node.get_text("\n", strip=True)
            if len(text) > 180:
                candidates.append(text)

    if not candidates:
        body = soup.body
        if body:
            candidates.append(body.get_text("\n", strip=True))

    if not candidates:
        return None

    text = max(candidates, key=len)
    text = clean_text((meta_desc + "\n" + text).strip())

    if len(text) < 150:
        return None

    if not is_relevant_oyster(title + " " + text):
        return None

    return {
        "title": title[:180] or url,
        "content": text,
        "url": url,
        "date": extract_date(text + " " + url),
        "source": get_domain(url)
    }


def crawl_public_web_pages():
    urls = collect_web_urls()
    docs = []

    print("📄 开始抓取公开网页正文...")
    for idx, url in enumerate(urls, start=1):
        print(f"  [{idx}/{len(urls)}] {url}")
        item = extract_page_content(url)
        if not item:
            continue

        chunks = chunk_text(item["content"])
        for i, chunk in enumerate(chunks, start=1):
            title = item["title"] if len(chunks) == 1 else f"{item['title']} - 片段{i}"
            docs.append(create_doc(
                title=title,
                content=chunk,
                source=f"公开网页：{item['source']}",
                url=item["url"],
                date=item["date"],
                chunk_index=i,
                source_type="web"
            ))

        time.sleep(REQUEST_SLEEP)

    print(f"  ✅ 公开网页生成知识片段 {len(docs)} 条")
    return docs


# ---------------------------------------------------------
# Crossref
# ---------------------------------------------------------
def crawl_crossref():
    print("📚 抓取 Crossref 公开论文题录/摘要...")
    docs = []

    queries = ENGLISH_KEYWORDS + [
        "oyster aquaculture China",
        "Crassostrea gigas triploid",
        "oyster breeding disease mortality",
        "Pacific oyster genetic resources"
    ]

    for q in queries:
        params = {
            "query": q,
            "rows": 10,
            "select": "title,abstract,DOI,published-print,published-online,container-title,URL"
        }

        resp = safe_get("https://api.crossref.org/works", params=params)
        if not resp:
            continue

        try:
            data = resp.json()
            items = data.get("message", {}).get("items", [])
        except Exception:
            items = []

        for item in items:
            title_list = item.get("title") or []
            title = title_list[0] if title_list else q

            abstract = item.get("abstract", "")
            abstract = re.sub(r"<[^>]+>", "", abstract)
            abstract = clean_text(abstract)

            container = ""
            if item.get("container-title"):
                container = item["container-title"][0]

            doi = item.get("DOI", "")
            item_url = item.get("URL", "")

            if not abstract:
                abstract = f"题名：{title}。期刊/来源：{container}。DOI：{doi}。该记录来自 Crossref 开放题录，未提供摘要全文。"

            if not is_relevant_oyster(title + " " + abstract):
                continue

            content = f"标题：{title}\n期刊/来源：{container}\nDOI：{doi}\n摘要/题录：{abstract}"

            docs.append(create_doc(
                title=title,
                content=content,
                source="Crossref公开论文题录",
                url=item_url,
                date="",
                chunk_index=1,
                source_type="paper"
            ))

        time.sleep(REQUEST_SLEEP)

    print(f"  ✅ Crossref 生成知识片段 {len(docs)} 条")
    return docs


# ---------------------------------------------------------
# Semantic Scholar
# ---------------------------------------------------------
def crawl_semantic_scholar():
    print("📚 抓取 Semantic Scholar 公开论文题录/摘要...")
    docs = []

    queries = ENGLISH_KEYWORDS + [
        "Crassostrea gigas oyster aquaculture",
        "triploid oyster Crassostrea gigas",
        "Pacific oyster disease mortality",
        "Pacific oyster genetic resources"
    ]

    for q in queries:
        params = {
            "query": q,
            "limit": 10,
            "fields": "title,abstract,year,authors,url,venue"
        }

        resp = safe_get("https://api.semanticscholar.org/graph/v1/paper/search", params=params)
        if not resp:
            continue

        try:
            data = resp.json()
            items = data.get("data", [])
        except Exception:
            items = []

        for item in items:
            title = item.get("title", q)
            abstract = clean_text(item.get("abstract", "") or "")
            year = str(item.get("year") or "")
            venue = item.get("venue") or ""
            item_url = item.get("url") or ""

            authors = item.get("authors") or []
            author_names = ", ".join([a.get("name", "") for a in authors[:5]])

            if not abstract:
                abstract = f"题名：{title}。年份：{year}。来源：{venue}。作者：{author_names}。该记录来自 Semantic Scholar 开放题录，未提供摘要全文。"

            if not is_relevant_oyster(title + " " + abstract):
                continue

            content = f"标题：{title}\n作者：{author_names}\n年份：{year}\n来源：{venue}\n摘要/题录：{abstract}"

            docs.append(create_doc(
                title=title,
                content=content,
                source="Semantic Scholar公开论文题录",
                url=item_url,
                date=year,
                chunk_index=1,
                source_type="paper"
            ))

        time.sleep(REQUEST_SLEEP)

    print(f"  ✅ Semantic Scholar 生成知识片段 {len(docs)} 条")
    return docs


# ---------------------------------------------------------
# SerpAPI：谷歌学术/知网/百度学术等结果扩展
# ---------------------------------------------------------
def serpapi_get(params):
    if not SERPAPI_KEY:
        return None
    params = dict(params)
    params["api_key"] = SERPAPI_KEY
    resp = safe_get("https://serpapi.com/search.json", params=params, timeout=30)
    if not resp:
        return None
    try:
        return resp.json()
    except Exception:
        return None


def crawl_google_scholar_via_serpapi():
    print("🎓 SerpAPI 谷歌学术检索...")
    docs = []

    if not SERPAPI_KEY:
        print("  ℹ️ 未配置 SERPAPI_KEY，跳过谷歌学术增强检索")
        return docs

    for q in SCHOLAR_QUERIES:
        data = serpapi_get({
            "engine": "google_scholar",
            "q": q,
            "hl": "zh-cn",
            "num": MAX_SERP_RESULTS_PER_QUERY
        })
        if not data:
            continue

        for item in data.get("organic_results", []):
            title = item.get("title", "")
            link = item.get("link", "")
            snippet = item.get("snippet", "")
            publication = item.get("publication_info", {}).get("summary", "")

            if not title or not is_relevant_oyster(title + " " + snippet):
                continue

            content = (
                f"标题：{title}\n"
                f"来源信息：{publication}\n"
                f"摘要/片段：{snippet}\n"
                f"链接：{link}"
            )

            docs.append(create_doc(
                title=title,
                content=content,
                source="SerpAPI谷歌学术题录/摘要",
                url=link,
                date=extract_date(publication + " " + snippet),
                source_type="scholar"
            ))

        time.sleep(REQUEST_SLEEP)

    print(f"  ✅ 谷歌学术生成知识片段 {len(docs)} 条")
    return docs


def crawl_cnki_baidu_wanfang_via_serpapi():
    print("🎓 SerpAPI 中文学术来源检索：知网/百度学术/万方/维普...")
    docs = []

    if not SERPAPI_KEY:
        print("  ℹ️ 未配置 SERPAPI_KEY，跳过中文学术增强检索")
        return docs

    site_queries = []
    for q in SCHOLAR_QUERIES:
        site_queries.extend([
            f"{q} site:kns.cnki.net",
            f"{q} site:cnki.net",
            f"{q} site:xueshu.baidu.com",
            f"{q} site:wanfangdata.com.cn",
            f"{q} site:cqvip.com",
        ])

    # 控制总量，避免 SerpAPI 消耗过快
    site_queries = site_queries[:30]

    for q in site_queries:
        data = serpapi_get({
            "engine": "google",
            "q": q,
            "hl": "zh-cn",
            "num": MAX_SERP_RESULTS_PER_QUERY
        })
        if not data:
            continue

        for item in data.get("organic_results", []):
            title = item.get("title", "")
            link = item.get("link", "")
            snippet = item.get("snippet", "")

            if not title or not link:
                continue
            if not is_relevant_oyster(title + " " + snippet):
                continue

            domain = get_domain(link)
            source_name = "中文学术检索结果"
            if "cnki" in domain:
                source_name = "知网公开检索题录/摘要"
            elif "xueshu.baidu" in domain:
                source_name = "百度学术公开检索题录/摘要"
            elif "wanfangdata" in domain:
                source_name = "万方公开检索题录/摘要"
            elif "cqvip" in domain:
                source_name = "维普公开检索题录/摘要"

            content = (
                f"标题：{title}\n"
                f"来源域名：{domain}\n"
                f"摘要/搜索片段：{snippet}\n"
                f"链接：{link}\n"
                f"说明：该记录来自搜索结果公开题录/摘要片段，不代表已抓取论文全文。"
            )

            docs.append(create_doc(
                title=title,
                content=content,
                source=source_name,
                url=link,
                date=extract_date(snippet),
                source_type="scholar"
            ))

        time.sleep(REQUEST_SLEEP)

    print(f"  ✅ 中文学术检索生成知识片段 {len(docs)} 条")
    return docs


# ---------------------------------------------------------
# data/raw 解析
# ---------------------------------------------------------
def read_pdf(path):
    items = []
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                text = clean_text(text)
                if text:
                    items.append((text, f"第{i}页"))
    except Exception as e:
        print(f"  ⚠️ PDF解析失败：{path.name}，原因：{e}")
    return items


def read_docx(path):
    items = []
    try:
        import docx
        doc = docx.Document(str(path))
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = clean_text("\n".join(paragraphs))
        if text:
            items.append((text, "全文"))
    except Exception as e:
        print(f"  ⚠️ DOCX解析失败：{path.name}，原因：{e}")
    return items


def read_txt_md(path):
    items = []
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
        text = clean_text(text)
        if text:
            items.append((text, "全文"))
    except Exception as e:
        print(f"  ⚠️ 文本解析失败：{path.name}，原因：{e}")
    return items


def read_csv_file(path):
    items = []
    try:
        rows_text = []
        with open(path, "r", encoding="utf-8-sig", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            for idx, row in enumerate(reader, start=1):
                row_text = " | ".join([cell.strip() for cell in row if cell.strip()])
                if row_text:
                    rows_text.append(f"第{idx}行：{row_text}")
        text = clean_text("\n".join(rows_text))
        if text:
            items.append((text, "CSV全文"))
    except Exception as e:
        print(f"  ⚠️ CSV解析失败：{path.name}，原因：{e}")
    return items


def read_xlsx_file(path):
    items = []
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True)
        for sheet in wb.worksheets:
            rows_text = []
            for idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if values:
                    rows_text.append(f"第{idx}行：{' | '.join(values)}")
            text = clean_text("\n".join(rows_text))
            if text:
                items.append((text, f"工作表：{sheet.title}"))
    except Exception as e:
        print(f"  ⚠️ Excel解析失败：{path.name}，原因：{e}")
    return items


def read_file(path):
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix in [".txt", ".md"]:
        return read_txt_md(path)
    if suffix == ".csv":
        return read_csv_file(path)
    if suffix == ".xlsx":
        return read_xlsx_file(path)
    return []


def build_docs_from_raw_files():
    print("📁 检查 data/raw 本地资料...")
    docs = []

    if not RAW_DIR.exists():
        print("  ℹ️ 未发现 data/raw，跳过本地资料解析")
        return docs

    files = [
        p for p in RAW_DIR.rglob("*")
        if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES
    ]

    if not files:
        print("  ℹ️ data/raw 没有可解析文件")
        return docs

    print(f"  ✅ 发现本地资料 {len(files)} 个")

    for path in files:
        print(f"  📄 解析：{path}")
        items = read_file(path)

        for text, locator in items:
            chunks = chunk_text(text)
            for i, chunk in enumerate(chunks, start=1):
                title = f"{path.stem} - 片段{i}"
                docs.append(create_doc(
                    title=title,
                    content=chunk,
                    source=f"本地资料：{path.as_posix()} {locator}",
                    url="",
                    file_name=path.name,
                    chunk_index=i,
                    source_type="local"
                ))

    print(f"  ✅ 本地资料生成知识片段 {len(docs)} 条")
    return docs


# ---------------------------------------------------------
# 清洗、去重、排序、限制体积
# ---------------------------------------------------------
def remove_known_bad_records(docs):
    """
    清除已知错误污染项。
    例如 GB/T 24860-2010 实为圆斑星鲽，不得作为太平洋牡蛎标准。
    """
    cleaned = []
    for d in docs:
        text = f"{d.get('title','')} {d.get('content','')}"
        if "GB/T 24860" in text and ("太平洋牡蛎" in text or "牡蛎" in text):
            print(f"  🧹 删除疑似错误标准记录：{d.get('title')}")
            continue
        cleaned.append(d)
    return cleaned


def deduplicate_docs(docs):
    seen = set()
    unique = []

    for doc in docs:
        title = doc.get("title", "")
        content = doc.get("content", "")
        url = doc.get("url", "")

        key = make_id((url or title) + content[:350])
        if key in seen:
            continue

        seen.add(key)
        unique.append(doc)

    return unique


def limit_docs(docs, max_docs=500):
    def score(doc):
        s = 0
        src = doc.get("source", "")
        url = doc.get("url", "")
        domain = doc.get("domain", "")

        if doc.get("has_url"):
            s += 25
        if "本地资料" in src:
            s += 20
        for d in PREFERRED_DOMAINS:
            if d in url or d in src or d in domain:
                s += 25
        if "谷歌学术" in src or "知网" in src or "百度学术" in src:
            s += 20
        if "Crossref" in src or "Semantic" in src:
            s += 15
        s += min(doc.get("content_length", 0) / 100, 25)
        return s

    docs.sort(key=score, reverse=True)
    return docs[:max_docs]


def main():
    print("=" * 70)
    print("🦪 生蚝AI知识库 - 去兜底 + 扩大检索来源构建")
    print(f"⏰ 时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 70)

    docs = []

    try:
        docs.extend(crawl_public_web_pages())
    except Exception as e:
        print(f"⚠️ 公开网页抓取整体异常：{e}")

    try:
        docs.extend(crawl_crossref())
    except Exception as e:
        print(f"⚠️ Crossref 抓取异常：{e}")

    try:
        docs.extend(crawl_semantic_scholar())
    except Exception as e:
        print(f"⚠️ Semantic Scholar 抓取异常：{e}")

    try:
        docs.extend(crawl_google_scholar_via_serpapi())
    except Exception as e:
        print(f"⚠️ 谷歌学术 SerpAPI 抓取异常：{e}")

    try:
        docs.extend(crawl_cnki_baidu_wanfang_via_serpapi())
    except Exception as e:
        print(f"⚠️ 中文学术 SerpAPI 抓取异常：{e}")

    try:
        docs.extend(build_docs_from_raw_files())
    except Exception as e:
        print(f"⚠️ 本地资料解析异常：{e}")

    docs = remove_known_bad_records(docs)
    docs = deduplicate_docs(docs)
    docs = limit_docs(docs, max_docs=500)

    category_count = {}
    source_count = {}
    url_count = 0

    for doc in docs:
        category = doc.get("category", "未分类")
        source = doc.get("source", "未知来源")
        category_count[category] = category_count.get(category, 0) + 1
        source_count[source] = source_count.get(source, 0) + 1
        if doc.get("has_url"):
            url_count += 1

    output = {
        "name": "生蚝AI知识库",
        "version": "4.0-no-builtin-expanded-sources",
        "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total": len(docs),
        "url_count": url_count,
        "categories": category_count,
        "sources": source_count,
        "documents": docs
    }

    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print("=" * 70)
    print(f"✅ 构建完成，共生成 {len(docs)} 条知识片段")
    print(f"🔗 带 URL 的资料：{url_count} 条")
    print("📊 分类统计：")
    for k, v in category_count.items():
        print(f"  - {k}: {v}")
    print(f"📁 输出文件：{OUTPUT_FILE}")
    print("=" * 70)


if __name__ == "__main__":
    main()
